"""
dashboard.py — Streamlit QA Command Center dashboard.

Run with:
    python -m streamlit run dashboard.py
"""
import json
import io
import tempfile
import os
import pandas as pd
import streamlit as st
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

from db import get_connection, init_db, save_evaluation
from engine import QAEngine
from file_parser import parse_filename

# ---------------------------------------------------------------------------
# Word transcript export helper
# ---------------------------------------------------------------------------
def build_transcript_docx(rows: pd.DataFrame) -> bytes:
    doc = Document()

    # Document title
    title = doc.add_heading("QA Transcript Report", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    from datetime import datetime
    sub = doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub.runs[0].font.color.rgb = RGBColor(0x9C, 0xA3, 0xAF)
    doc.add_paragraph()

    for _, row in rows.iterrows():
        score = row.get("qa_score") or 0
        agent   = str(row.get("agent_name") or "Unknown")
        call_id = str(row.get("call_uuid") or "N/A")

        # --- Call header ---
        heading = doc.add_heading(f"📞  {call_id}", level=1)
        heading.runs[0].font.size = Pt(13)

        # Score color indicator
        if score >= 80:
            label, color = "PASS", RGBColor(0x10, 0xB9, 0x81)
        elif score < 60:
            label, color = "FAIL", RGBColor(0xEF, 0x44, 0x44)
        else:
            label, color = "AVERAGE", RGBColor(0xF5, 0x9E, 0x0B)

        score_para = doc.add_paragraph()
        score_run = score_para.add_run(f"Score: {score}/100  [{label}]")
        score_run.bold = True
        score_run.font.color.rgb = color
        score_run.font.size = Pt(11)

        # Meta table
        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        hdr = table.rows[0].cells
        hdr[0].text = f"Agent\n{agent}"
        hdr[1].text = f"Date\n{str(row.get('call_date') or 'N/A')}"
        hdr[2].text = f"Loan Type\n{str(row.get('detected_loan_type') or 'N/A')}"
        hdr[3].text = f"Category\n{str(row.get('category') or 'N/A')}"
        for cell in hdr:
            for para in cell.paragraphs:
                for run in para.runs:
                    run.font.size = Pt(9)

        doc.add_paragraph()

        # AI summary & feedback
        raw_json = row.get("grading_json")
        try:
            analysis = json.loads(raw_json) if isinstance(raw_json, str) else (raw_json or {})
        except Exception:
            analysis = {}

        summary  = str(analysis.get("summary")  or row.get("qa_summary")  or "")
        feedback = str(analysis.get("areas_of_improvement") or row.get("qa_feedback") or "")
        issue    = str(analysis.get("detected_issue") or "None")
        reasoning = str(analysis.get("reasoning") or "")

        if summary:
            p = doc.add_paragraph()
            p.add_run("Summary: ").bold = True
            p.add_run(summary)

        if feedback:
            p = doc.add_paragraph()
            p.add_run("Areas of Improvement: ").bold = True
            p.add_run(feedback)

        if issue and issue.lower() != "none":
            p = doc.add_paragraph()
            p.add_run("Detected Issue: ").bold = True
            p.add_run(issue)

        if reasoning:
            p = doc.add_paragraph()
            p.add_run("Reasoning: ").bold = True
            p.add_run(reasoning)

        doc.add_paragraph()

        # Transcript
        transcript = (row.get("transcription_text") or "").strip()
        if transcript:
            t_heading = doc.add_heading("Transcript", level=2)
            t_heading.runs[0].font.size = Pt(11)
            for line in transcript.split("\n"):
                line = line.strip()
                if not line:
                    continue
                p = doc.add_paragraph(style="Normal")
                if line.startswith("Speaker "):
                    colon = line.find(":")
                    if colon != -1:
                        speaker_run = p.add_run(line[:colon + 1] + " ")
                        speaker_run.bold = True
                        speaker_run.font.color.rgb = RGBColor(0x3B, 0x82, 0xF6)
                        p.add_run(line[colon + 1:].strip())
                    else:
                        p.add_run(line)
                else:
                    p.add_run(line)
                p.paragraph_format.space_after = Pt(2)

        # Page break between calls (except last)
        if _ != rows.index[-1]:
            doc.add_page_break()

    buf = io.BytesIO()
    doc.save(buf)
    return buf.getvalue()


# ---------------------------------------------------------------------------
# Page setup
# ---------------------------------------------------------------------------
st.set_page_config(page_title="QA Command Center", layout="wide")

st.markdown("""
    <style>
    .main { background-color: #0e1117; }
    .stMetric {
        background-color: #1f2937;
        padding: 15px;
        border-radius: 10px;
        border: 1px solid #374151;
    }
    .call-card {
        background-color: #1f2937;
        padding: 20px;
        border-radius: 10px;
        border-left: 5px solid #3b82f6;
        margin-bottom: 10px;
    }
    .fail-card { border-left: 5px solid #ef4444 !important; }
    .pass-card { border-left: 5px solid #10b981 !important; }
    </style>
""", unsafe_allow_html=True)


# ---------------------------------------------------------------------------
# Database helpers
# ---------------------------------------------------------------------------
@st.cache_data(ttl=30)
def load_data() -> pd.DataFrame:
    init_db()
    conn = get_connection()
    df   = pd.read_sql_query(
        "SELECT * FROM CallEvaluations ORDER BY call_timestamp DESC", conn
    )
    conn.close()
    return df


# ---------------------------------------------------------------------------
# Header
# ---------------------------------------------------------------------------
st.title("🛡️ QA Command Center")
st.markdown("### Real-time AI Call Auditing & Loan Verification")

# ---------------------------------------------------------------------------
# Upload Section
# ---------------------------------------------------------------------------
with st.expander("📁 Upload & Process Audio Files", expanded=True):
    uploaded_files = st.file_uploader(
        "Drop MP3 or WAV files here",
        type=["mp3", "wav"],
        accept_multiple_files=True,
    )

    if uploaded_files:
        st.markdown("**Optional metadata** (applies to all uploaded files if filled in)")
        col1, col2, col3, col4 = st.columns(4)
        with col1:
            agent_name = st.text_input("Agent Name")
        with col2:
            call_date = st.date_input("Call Date")
        with col3:
            duration = st.number_input("Duration (seconds)", min_value=0, value=0)
        with col4:
            hangup = st.selectbox("Hangup Source", ["", "Agent", "Customer", "System"])

        transfer = st.text_input("Transfer Destination (optional)")

        if st.button("🚀 Process Files", type="primary"):
            engine = QAEngine()
            conn   = get_connection()

            for uploaded_file in uploaded_files:
                filename = uploaded_file.name
                st.write(f"**Processing:** `{filename}`")

                # Write to a temp file so Deepgram can read it
                suffix = os.path.splitext(filename)[1]
                with tempfile.NamedTemporaryFile(delete=False, suffix=suffix) as tmp:
                    tmp.write(uploaded_file.read())
                    tmp_path = tmp.name

                try:
                    with st.spinner(f"Transcribing {filename}..."):
                        transcript = engine.transcribe_audio(tmp_path)

                    # Auto-parse metadata from filename, let manual inputs override
                    parsed = parse_filename(filename)
                    metadata = {
                        "agent_name":   agent_name or parsed.get("agent_name"),
                        "agent_username": parsed.get("agent_username"),
                        "loan_hint": parsed.get("loan_hint"),
                        "lead_phone":   parsed.get("lead_phone"),
                        "call_date":    str(call_date) if call_date else parsed.get("call_date"),
                        "duration":     duration or None,
                        "hangup":       hangup or None,
                        "transfer_ext": transfer or None,
                    }

                    with st.spinner(f"Grading {filename} with AI..."):
                        analysis = engine.evaluate_call(transcript, metadata)

                    # Use AI-detected agent name if we still don't have one
                    if not metadata["agent_name"] and analysis.get("agent_name"):
                        metadata["agent_name"] = analysis["agent_name"]

                    save_evaluation(conn, filename, metadata, transcript, analysis)

                    score = analysis.get("score", 0)
                    color = "green" if score >= 80 else "red" if score < 60 else "orange"
                    st.success(
                        f"✓ **{filename}** — Score: :{color}[**{score}/100**] | "
                        f"Loan: {analysis.get('loan_type', 'Unknown')} | "
                        f"Category: {analysis.get('category', 'N/A')}"
                    )

                except Exception as e:
                    st.error(f"✗ Error processing {filename}: {e}")

                finally:
                    os.unlink(tmp_path)

            conn.close()
            st.cache_data.clear()
            st.rerun()

# ---------------------------------------------------------------------------
# Load data
# ---------------------------------------------------------------------------
df = load_data()

if df.empty:
    st.warning(
        "No evaluations yet. Upload audio files above, or drop .txt transcripts "
        "into the `output/` folder and run `python critiera.py`."
    )
    st.stop()

# ---------------------------------------------------------------------------
# Top-level metrics
# ---------------------------------------------------------------------------
m1, m2, m3, m4 = st.columns(4)
avg_score = round(df["qa_score"].dropna().mean(), 1) if not df["qa_score"].isna().all() else 0
m1.metric("Average QA Score", f"{avg_score}%")
m2.metric("Total Audits",     len(df))
m3.metric("Critical Flags",   int((df["qa_score"] < 60).sum()))
m4.metric("Pass Rate (≥80)",  f"{round((df['qa_score'] >= 80).mean() * 100)}%")

# ---------------------------------------------------------------------------
# Sidebar filters
# ---------------------------------------------------------------------------
st.sidebar.title("🔍 Search & Filter")
search_uuid   = st.sidebar.text_input("Search by Call ID / Filename")
agent_options = sorted(df["agent_name"].dropna().unique().tolist())
agent_filter  = st.sidebar.multiselect("Filter by Agent", options=agent_options)
loan_options  = sorted(df["detected_loan_type"].dropna().unique().tolist())
loan_filter   = st.sidebar.multiselect("Filter by Loan Type", options=loan_options)
cat_options   = sorted(df["category"].dropna().unique().tolist())
cat_filter    = st.sidebar.multiselect("Filter by Category", options=cat_options)

if st.sidebar.button("🔄 Refresh Data"):
    st.cache_data.clear()
    st.rerun()

# ---------------------------------------------------------------------------
# Filter logic
# ---------------------------------------------------------------------------
filtered = df.copy()
if agent_filter:
    filtered = filtered[filtered["agent_name"].isin(agent_filter)]
if loan_filter:
    filtered = filtered[filtered["detected_loan_type"].isin(loan_filter)]
if cat_filter:
    filtered = filtered[filtered["category"].isin(cat_filter)]
if search_uuid:
    filtered = filtered[
        filtered["call_uuid"].str.contains(search_uuid, case=False, na=False)
    ]

# ---------------------------------------------------------------------------
# Tabs
# ---------------------------------------------------------------------------
tab1, tab2 = st.tabs(["📊 Performance Overview", "🎧 Audit Queue"])

# --- Tab 1: Charts ---
with tab1:
    col_a, col_b = st.columns(2)

    with col_a:
        st.subheader("Agent Average Score")
        chart_data = (
            filtered.groupby("agent_name")["qa_score"]
            .mean()
            .sort_values(ascending=False)
        )
        st.bar_chart(chart_data)

    with col_b:
        st.subheader("Score Distribution")
        score_bins = pd.cut(
            filtered["qa_score"].dropna(),
            bins=[0, 59, 79, 100],
            labels=["Fail (<60)", "Average (60-79)", "Pass (≥80)"]
        ).value_counts()
        st.bar_chart(score_bins)

    st.subheader("All Evaluations")
    display_cols = [
        "call_uuid", "agent_name", "call_date", "category",
        "detected_loan_type", "qa_score", "call_timestamp"
    ]
    available = [c for c in display_cols if c in filtered.columns]
    st.dataframe(filtered[available], use_container_width=True)

# --- Tab 2: Audit cards ---
with tab2:
    # Export buttons
    col_title, col_csv, col_docx = st.columns([3, 1, 1])
    with col_title:
        st.subheader(f"Showing {len(filtered)} evaluations")
    with col_csv:
        # Build export dataframe with all useful fields
        export_cols = [
            "call_uuid", "agent_name", "lead_phone", "call_date",
            "call_timestamp", "duration_seconds", "hangup_source",
            "transfer_destination", "detected_loan_type", "category",
            "qa_score", "qa_summary", "qa_feedback",
        ]
        export_df = filtered[[c for c in export_cols if c in filtered.columns]].copy()
        export_df.rename(columns={
            "call_uuid":            "File / Call ID",
            "agent_name":           "Agent Name",
            "lead_phone":           "Customer Phone",
            "call_date":            "Call Date",
            "call_timestamp":       "Processed At",
            "duration_seconds":     "Duration (s)",
            "hangup_source":        "Hangup Source",
            "transfer_destination": "Transfer Destination",
            "detected_loan_type":   "Loan Type",
            "category":             "Category",
            "qa_score":             "QA Score",
            "qa_summary":           "Summary",
            "qa_feedback":          "Areas of Improvement",
        }, inplace=True)

        st.download_button(
            label="⬇️ Export CSV",
            data=export_df.to_csv(index=False).encode("utf-8"),
            file_name=f"audit_queue_export.csv",
            mime="text/csv",
            use_container_width=True,
        )

    with col_docx:
        docx_bytes = build_transcript_docx(filtered)
        st.download_button(
            label="📄 Export Transcripts",
            data=docx_bytes,
            file_name="transcripts_export.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            use_container_width=True,
        )

    for _, row in filtered.iterrows():
        score = row.get("qa_score") or 0
        if score >= 80:
            card_style = "pass-card"
        elif score < 60:
            card_style = "fail-card"
        else:
            card_style = ""

        score_color = "#10b981" if score >= 80 else "#ef4444" if score < 60 else "#f59e0b"

        st.markdown(f"""
            <div class="call-card {card_style}">
                <div style="display:flex; justify-content:space-between;">
                    <span style="font-size:1.1rem; font-weight:bold;">
                        {row.get('call_uuid', 'N/A')}
                    </span>
                    <span style="color:#9ca3af;">{row.get('call_timestamp', '')}</span>
                </div>
                <div style="display:flex; gap:20px; margin-top:10px; flex-wrap:wrap;">
                    <div><b>Agent:</b> {row.get('agent_name', 'N/A')}</div>
                    <div><b>Loan:</b> {row.get('detected_loan_type', 'N/A')}</div>
                    <div><b>Category:</b> {row.get('category', 'N/A')}</div>
                    <div><b>Duration:</b> {row.get('duration_seconds') or 'N/A'}s</div>
                    <div style="color:{score_color}; font-weight:bold;">
                        Score: {score}%
                    </div>
                </div>
            </div>
        """, unsafe_allow_html=True)

        with st.expander("View Full Audit Details & Transcript"):
            col_left, col_right = st.columns([1, 2])

            raw_json = row.get("grading_json")
            try:
                analysis = json.loads(raw_json) if isinstance(raw_json, str) else (raw_json or {})
            except (json.JSONDecodeError, TypeError):
                analysis = {}

            with col_left:
                st.write("### AI Verdict")
                st.write(f"**Issue:** {analysis.get('detected_issue', 'None')}")
                st.write(f"**Transfer Goal:** {row.get('transfer_destination', 'N/A')}")
                st.progress(int(score) / 100)
                st.caption(analysis.get("summary") or row.get("qa_summary") or "")
                if analysis.get("areas_of_improvement") or row.get("qa_feedback"):
                    st.write("**Feedback:**")
                    st.info(analysis.get("areas_of_improvement") or row.get("qa_feedback"))

            with col_right:
                st.write("### Call Transcript")
                transcript = row.get("transcription_text") or ""
                st.text_area(
                    "Dialogue",
                    transcript,
                    height=250,
                    key=f"tx_{row.get('call_uuid', _)}",
                )
